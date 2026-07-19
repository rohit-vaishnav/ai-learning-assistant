import docx

def read_docx(file_path: str) -> list[dict]:
    """
    Reads a DOCX file and extracts text paragraph-by-paragraph.
    Returns:
        list of dict: [{"text": full_text, "page": 1}] (docx files are unstructured, so page defaults to 1)
    """
    try:
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            # Try to read tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            full_text = "\n".join(text_parts)

        return [{"text": full_text, "page": 1}]
    except Exception as e:
        raise RuntimeError(f"Error parsing DOCX: {str(e)}")
